from fastapi import FastAPI, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from docx import Document
import pdfplumber
import tempfile
import os
import uuid
import requests

# ======================================================
# APP
# ======================================================
app = FastAPI(title="AI Resume Analyzer")

# ✅ HEALTH CHECK
@app.get("/")
def health():
    return {"status": "Backend running"}

# ======================================================
# CORS
# ======================================================
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],   # OK for now (local + deployment)
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ======================================================
# OPENROUTER CONFIG (SAFE – NO CRASH)
# ======================================================
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"
OPENROUTER_MODEL = "meta-llama/llama-3.1-8b-instruct"

if not OPENROUTER_API_KEY:
    print("WARNING: OPENROUTER_API_KEY not set. AI analysis will be disabled.")

# ======================================================
# SKILLS
# ======================================================
COMMON_SKILLS = [
    "python", "java", "javascript", "react", "next.js",
    "node.js", "fastapi", "sql", "mysql", "postgresql",
    "mongodb", "aws", "docker", "kubernetes",
    "git", "github", "rest api", "html", "css",
    "flutter", "dart",
    "machine learning", "data analysis"
]

# ======================================================
# TEXT EXTRACTION
# ======================================================
def extract_text_from_pdf(path: str) -> str:
    text = ""
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() or ""
    return text.strip()

def extract_text_from_docx(path: str) -> str:
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs).strip()

def extract_skills(text: str):
    t = text.lower()
    return [s for s in COMMON_SKILLS if s in t]

# ======================================================
# ATS SCORE
# ======================================================
def ats_keyword_score(resume_text: str, jd_skills):
    if not jd_skills:
        return 0
    resume_text = resume_text.lower()
    matched = sum(1 for s in jd_skills if s in resume_text)
    return int((matched / len(jd_skills)) * 100)

# ======================================================
# OPENROUTER AI CALL (SAFE)
# ======================================================
def openrouter_think(prompt: str) -> str:
    if not OPENROUTER_API_KEY:
        return "AI analysis unavailable (API key not configured)."

    safe_prompt = prompt[:3000]

    try:
        response = requests.post(
            OPENROUTER_URL,
            headers={
                "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                "Content-Type": "application/json",
                "HTTP-Referer": "http://localhost",
                "X-Title": "AI Resume Improver"
            },
            json={
                "model": OPENROUTER_MODEL,
                "messages": [
                    {
                        "role": "system",
                        "content": "You are a senior hiring manager. Be realistic, strict, and practical."
                    },
                    {
                        "role": "user",
                        "content": safe_prompt
                    }
                ],
                "temperature": 0.4,
                "max_tokens": 700
            },
            timeout=60
        )

        if response.status_code != 200:
            return f"OpenRouter error: {response.text}"

        return response.json()["choices"][0]["message"]["content"].strip()

    except Exception as e:
        return f"AI error: {str(e)}"

# ======================================================
# AI ANALYSIS
# ======================================================
def ai_analysis(resume_text, job_description, score, matched, missing):
    prompt = f"""
Job Description:
{job_description[:800]}

Resume:
{resume_text[:800]}

ATS Score: {score}%

Matched Skills: {', '.join(matched) or 'None'}
Missing Skills: {', '.join(missing) or 'None'}

Answer as a hiring manager:
1. Why is the ATS score this value?
2. Resume weaknesses.
3. Skills to prioritize.
4. Real projects to build for this role.
5. A realistic 60-day improvement plan.
"""
    return {
        "analysis_text": openrouter_think(prompt)
    }

# ======================================================
# DOCX UPDATE (ONLY FOR DOCX)
# ======================================================
def update_docx_resume(path, summary):
    doc = Document(path)

    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().upper() == "SUMMARY" and i + 1 < len(doc.paragraphs):
            doc.paragraphs[i + 1].text = summary
            break

    os.makedirs("generated", exist_ok=True)
    file_id = str(uuid.uuid4())
    out_path = f"generated/{file_id}.docx"
    doc.save(out_path)
    return file_id

# ======================================================
# API
# ======================================================
@app.post("/analyze-resume")
async def analyze_resume(
    resume: UploadFile = File(...),
    job_description: str = Form(...)
):
    suffix = resume.filename.split(".")[-1].lower()
    if suffix not in ["pdf", "docx"]:
        return {"error": "Only PDF and DOCX supported"}

    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{suffix}") as tmp:
        tmp.write(await resume.read())
        path = tmp.name

    resume_text = (
        extract_text_from_pdf(path)
        if suffix == "pdf"
        else extract_text_from_docx(path)
    )

    resume_skills = extract_skills(resume_text)
    jd_skills = extract_skills(job_description)

    matched = sorted(set(resume_skills) & set(jd_skills))
    missing = sorted(set(jd_skills) - set(resume_skills))

    score = ats_keyword_score(resume_text, jd_skills)
    skill_scores = {s: (100 if s in matched else 0) for s in jd_skills}

    download_id = None
    if suffix == "docx":
        download_id = update_docx_resume(
            path,
            "Results-driven candidate aligned with job requirements."
        )

    ai_result = ai_analysis(
        resume_text,
        job_description,
        score,
        matched,
        missing
    )

    return {
        "score": score,
        "matched_skills": matched,
        "missing_skills": missing,
        "skill_scores": skill_scores,
        "ai_analysis": ai_result,
        "download_id": download_id
    }

# ======================================================
# DOWNLOAD
# ======================================================
@app.get("/download/{file_id}")
def download(file_id: str):
    return FileResponse(
        f"generated/{file_id}.docx",
        filename="Improved_Resume.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
